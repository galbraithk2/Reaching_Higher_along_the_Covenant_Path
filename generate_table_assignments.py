from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Table assignments: (table_number, assigned_name)
assignments = [
    (1, "Newdale Ward"),
    (2, "Newdale Ward"),
    (3, "Aspen Ward"),
    (4, "Sugar Ward"),
    (5, "Madalyn Hunt"),
    (6, "Moody Creek Ward Families"),
    (7, "Institute Representatives"),
    (8, "Johnathan & Keli Huskinson"),
    (9, "Stake Emergency Committee"),
    (10, "Stake Emergency Committee"),
    (11, "North Fork, South Fork, & Teton Rivers"),
    (12, "Robert & Sandra Myers"),
    (13, "Canyon Creek Youth"),
    (14, "Ponderosa Ward"),
    (15, "Garth & Ruth Miller\nJoseph & Debra Cherrington\nBryce & Sherry Holman"),
    (16, "Heritage Park, Grand Teton, & Canyon Creek"),
    (17, "Rozan & Jerry Miller"),
    (18, "North Fork Ward Families"),
    (19, "Teton Island & Newdale"),
]

def add_half_block(doc, table_num, name, add_divider=False):
    """Add a centered half-page block with Table # and name."""
    # Spacer before content to vertically center in half
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    # Set spacer height to push content toward center of its half
    pPr = spacer._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '1400')  # ~1 inch space before
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    # "Table X" line
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(f"Table {table_num}")
    run1.bold = True
    run1.font.size = Pt(48)
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(8)

    # Name line(s)
    for line in name.split('\n'):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(line)
        run2.bold = True
        run2.font.size = Pt(28)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)

    if add_divider:
        # Horizontal rule as a paragraph border
        div = doc.add_paragraph()
        div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr2 = div._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr2.append(pBdr)
        spacing2 = OxmlElement('w:spacing')
        spacing2.set(qn('w:before'), '600')
        spacing2.set(qn('w:after'), '0')
        pPr2.append(spacing2)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

doc = Document()

# Set page margins (0.75 inch all around)
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Default paragraph spacing
style = doc.styles['Normal']
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

# Group into pairs
pairs = []
for i in range(0, len(assignments), 2):
    pair = assignments[i:i+2]
    pairs.append(pair)

for page_idx, pair in enumerate(pairs):
    top = pair[0]
    add_half_block(doc, top[0], top[1], add_divider=True)

    if len(pair) == 2:
        bottom = pair[1]
        add_half_block(doc, bottom[0], bottom[1], add_divider=False)

    # Page break after each page except the last
    if page_idx < len(pairs) - 1:
        add_page_break(doc)

output_path = "/home/user/Reaching_Higher_along_the_Covenant_Path/Table_Assignments.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
